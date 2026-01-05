"""
Document Parser for ASHER
Supports: TXT, DOCX, MD, PDF, CSV, HTML, JSON
"""

import io
import json
import csv
from typing import Optional
from pypdf import PdfReader
from docx import Document
import markdown
from bs4 import BeautifulSoup


class DocumentParser:
    """Parse various document formats to extract text content"""

    @staticmethod
    def parse_txt(file_content: bytes) -> str:
        """Parse plain text file"""
        return file_content.decode('utf-8', errors='ignore')

    @staticmethod
    def parse_pdf(file_content: bytes) -> str:
        """Parse PDF file"""
        pdf_file = io.BytesIO(file_content)
        reader = PdfReader(pdf_file)

        text_parts = []
        for page_num, page in enumerate(reader.pages, 1):
            text = page.extract_text()
            if text.strip():
                text_parts.append(f"--- Page {page_num} ---\n{text}")

        return "\n\n".join(text_parts)

    @staticmethod
    def parse_docx(file_content: bytes) -> str:
        """Parse Word document"""
        docx_file = io.BytesIO(file_content)
        doc = Document(docx_file)

        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    text_parts.append(row_text)

        return "\n\n".join(text_parts)

    @staticmethod
    def parse_markdown(file_content: bytes) -> str:
        """Parse Markdown file"""
        md_text = file_content.decode('utf-8', errors='ignore')
        # Convert to HTML then extract text to handle formatting
        html = markdown.markdown(md_text)
        soup = BeautifulSoup(html, 'html.parser')
        return soup.get_text(separator='\n', strip=True)

    @staticmethod
    def parse_html(file_content: bytes) -> str:
        """Parse HTML file"""
        html_text = file_content.decode('utf-8', errors='ignore')
        soup = BeautifulSoup(html_text, 'html.parser')

        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.extract()

        return soup.get_text(separator='\n', strip=True)

    @staticmethod
    def parse_csv(file_content: bytes) -> str:
        """Parse CSV file"""
        csv_text = file_content.decode('utf-8', errors='ignore')
        csv_file = io.StringIO(csv_text)
        reader = csv.reader(csv_file)

        rows = []
        for row in reader:
            rows.append(" | ".join(row))

        return "\n".join(rows)

    @staticmethod
    def parse_json(file_content: bytes) -> str:
        """Parse JSON file"""
        json_text = file_content.decode('utf-8', errors='ignore')
        data = json.loads(json_text)
        # Pretty print JSON
        return json.dumps(data, indent=2)

    @classmethod
    def parse_file(cls, filename: str, file_content: bytes) -> str:
        """
        Parse a file based on its extension

        Args:
            filename: Name of the file
            file_content: Binary content of the file

        Returns:
            Extracted text content

        Raises:
            ValueError: If file type is not supported
        """
        ext = filename.lower().split('.')[-1]

        parsers = {
            'txt': cls.parse_txt,
            'pdf': cls.parse_pdf,
            'docx': cls.parse_docx,
            'doc': cls.parse_docx,
            'md': cls.parse_markdown,
            'markdown': cls.parse_markdown,
            'html': cls.parse_html,
            'htm': cls.parse_html,
            'csv': cls.parse_csv,
            'json': cls.parse_json,
            'jsonl': cls.parse_txt,  # JSONL is line-delimited JSON
        }

        parser = parsers.get(ext)
        if not parser:
            raise ValueError(f"Unsupported file type: .{ext}")

        try:
            return parser(file_content)
        except Exception as e:
            raise ValueError(f"Error parsing {ext} file: {str(e)}")

    @staticmethod
    def get_supported_extensions() -> list[str]:
        """Get list of supported file extensions"""
        return ['txt', 'pdf', 'docx', 'doc', 'md', 'markdown', 'html', 'htm', 'csv', 'json', 'jsonl']
